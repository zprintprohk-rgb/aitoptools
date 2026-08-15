# -*- coding: utf-8 -*-
"""渲染 Word 版《PHASE-2026-08-13-09-13 执行闭环报告》— 从定稿 MD 文档转换。
轻量 markdown->docx 转换：标题/表格/列表/粗体/引用/代码块。
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

WB = r'F:\aitoptools\.cluster\goal-flow-c52c6564-e8d7-44f3-a451-e84dc2949645'
DEST = os.path.join(WB, 'DELIVERY')
FILES = ['00-overview.md', '01-task-ledger.md', '02-artifacts.md',
         '03-risk-log.md', '04-review-record.md', '05-handoff.md']

INK = RGBColor(0x17, 0x20, 0x1c)
GREEN = RGBColor(0x0b, 0x5f, 0x59)
AMBER = RGBColor(0xd9, 0x77, 0x06)
MUTED = RGBColor(0x4a, 0x55, 0x68)

doc = Document()
# 页面与默认样式
for sec in doc.sections:
    sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.4)
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5


def add_heading(text, level):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = INK if level > 1 else GREEN
    if level == 1:
        run.font.size = Pt(16); run.font.bold = True
    elif level == 2:
        run.font.size = Pt(13); run.font.bold = True
    else:
        run.font.size = Pt(11.5); run.font.bold = True
    return h


def add_table(rows):
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = t.cell(i, j)
            txt = row[j] if j < len(row) else ''
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9)
            run.font.name = 'Microsoft YaHei'
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                from docx.oxml.ns import qn
                shd = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:fill'): '17201C'})
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


def convert_md(md_text):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1; continue
        # 表格
        if line.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells)
                i += 1
            add_table(rows)
            continue
        # 标题
        m = re.match(r'^(#{1,4})\s+(.*)$', line)
        if m:
            add_heading(m.group(2).strip(), min(len(m.group(1)), 4))
            i += 1; continue
        # 引用块
        if line.strip().startswith('>'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().lstrip('> '))
            run.font.color.rgb = MUTED
            run.font.size = Pt(9.5)
            p.paragraph_format.left_indent = Cm(0.6)
            i += 1; continue
        # 分隔线
        if re.match(r'^\s*[-*_]{3,}\s*$', line):
            i += 1; continue
        # 列表
        m = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(style='List Bullet' if m.group(2) in ('-', '*') else 'List Number')
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', m.group(3))
            text = re.sub(r'`(.*?)`', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
            i += 1; continue
        # 普通段落（处理粗体/行内码）
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        text = re.sub(r'`(.*?)`', r'\1', text)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        i += 1


# ============ 封面 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PHASE-2026-08-13-09-13 执行闭环报告')
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = INK
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('排名冲刺期 Phase 2 v2.0 · 三文件深度解读 → 台账 → 核验 → 修复 → 复核 → 交接')
run.font.size = Pt(11); run.font.color.rgb = GREEN
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('2026-08-15 · 主控汇总（5 路 Agent 独立核查 + 主线 git/cron 实证）· 功能主义网格版式')
run.font.size = Pt(9); run.font.color.rgb = MUTED
doc.add_paragraph()

# ============ 正文 ============
for f in FILES:
    path = os.path.join(DEST, f)
    if not os.path.exists(path):
        continue
    with open(path, encoding='utf-8') as fh:
        md = fh.read()
    # 跳过 md 自身的 H1 标题（文档名），用章节 H1 代替
    lines = md.splitlines()
    if lines and lines[0].startswith('# '):
        add_heading(lines[0][2:].split('—')[0].strip(), 1)
        convert_md('\n'.join(lines[1:]))
    else:
        convert_md(md)
    doc.add_paragraph()

# ============ 附录：修复三件套 ============
add_heading('附录：generate-sitemap.py 修复三件套（8/16 T2 消费）', 1)
app = doc.add_paragraph()
run = app.add_run('1. generate-sitemap_merged.py — 修复合并版脚本（标记保留/治愈 + 规范化去重 + robots 4 AI 代理 + heal 去重），整文件替换用；'
                  '2. sitemap_fix_v2.patch — LF 纯净 unified diff（6 hunks），git apply --check 已通过；'
                  '3. sitemap_fix_test_v2.py — 动态断言单测，27/27 PASS（当前生产 343 条基线）。'
                  '落地前 generate-pages.js 禁止运行；落地后首次运行「治愈标记」343→343 零增删。')
run.font.size = Pt(10)

out = os.path.join(DEST, 'PHASE-执行闭环报告.docx')
doc.save(out)
print('saved:', out, os.path.getsize(out), 'bytes')
